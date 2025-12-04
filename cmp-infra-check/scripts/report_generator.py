#!/usr/bin/env python3
"""
CMP 인프라 점검 보고서 생성 모듈
CSV 및 DOCX 형식 보고서 생성
"""

import csv
import os
from datetime import datetime
from typing import List, Dict, Any
from dataclasses import dataclass

# DOCX 생성
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ReportConfig:
    """보고서 설정"""
    report_type: str = "weekly"
    company_name: str = "CMP 인프라"
    team_name: str = "플랫폼팀"
    output_dir: str = "./output"


class CMPReportGenerator:
    """CMP 보고서 생성기"""
    
    def __init__(self, config: ReportConfig = None):
        self.config = config or ReportConfig()
        os.makedirs(self.config.output_dir, exist_ok=True)
    
    def _get_report_title(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"{now.year}년 {week_num}주차 CMP 인프라 정기점검 보고서"
        else:
            return f"{now.year}년 {now.month}월 CMP 인프라 정기점검 보고서"
    
    def _get_filename_prefix(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"cmp_infra_check_{now.year}_W{week_num:02d}"
        else:
            return f"cmp_infra_check_{now.year}_{now.month:02d}"
    
    def generate_csv(self, results: List[Dict], summary: Dict) -> str:
        """CSV 보고서 생성"""
        filename = f"{self._get_filename_prefix()}.csv"
        filepath = os.path.join(self.config.output_dir, filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            f.write(f"# {self._get_report_title()}\n")
            f.write(f"# 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# 회사: {self.config.company_name}\n")
            f.write(f"# 담당팀: {self.config.team_name}\n")
            f.write(f"# 총 점검항목: {summary.get('total', 0)}\n")
            f.write(f"# 정상: {summary.get('ok', 0)} / 경고: {summary.get('warning', 0)} / 위험: {summary.get('critical', 0)} / 확인불가: {summary.get('unknown', 0)}\n")
            f.write("\n")
            
            if results:
                writer = csv.DictWriter(f, fieldnames=results[0].keys())
                writer.writeheader()
                writer.writerows(results)
        
        return filepath
    
    def generate_docx(self, results: List[Dict], summary: Dict) -> str:
        """DOCX 보고서 생성"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 라이브러리가 설치되지 않았습니다.")
        
        filename = f"{self._get_filename_prefix()}.docx"
        filepath = os.path.join(self.config.output_dir, filename)
        
        doc = Document()
        
        # 제목
        title = doc.add_heading(self._get_report_title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 보고서 정보
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"회사: {self.config.company_name} | 담당팀: {self.config.team_name}")
        
        doc.add_paragraph()
        
        # 1. 요약 섹션
        doc.add_heading('1. 점검 결과 요약', level=1)
        
        summary_table = doc.add_table(rows=2, cols=5)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ['총 점검항목', '정상', '경고', '위험', '확인불가']
        hdr_cells = summary_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
        
        data_cells = summary_table.rows[1].cells
        data = [str(summary.get('total', 0)), str(summary.get('ok', 0)), 
                str(summary.get('warning', 0)), str(summary.get('critical', 0)), 
                str(summary.get('unknown', 0))]
        colors = [None, RGBColor(0, 128, 0), RGBColor(255, 165, 0), 
                  RGBColor(255, 0, 0), RGBColor(128, 128, 128)]
        
        for i, (value, color) in enumerate(zip(data, colors)):
            data_cells[i].text = value
            data_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if color:
                for run in data_cells[i].paragraphs[0].runs:
                    run.font.color.rgb = color
                    run.bold = True
        
        doc.add_paragraph()
        
        # 2. 환경별 점검 결과
        doc.add_heading('2. 환경별 점검 결과', level=1)
        
        by_env = summary.get('by_environment', {})
        if by_env:
            env_table = doc.add_table(rows=len(by_env) + 1, cols=5)
            env_table.style = 'Table Grid'
            
            env_headers = ['환경', '정상', '경고', '위험', '확인불가']
            for i, h in enumerate(env_headers):
                env_table.rows[0].cells[i].text = h
                env_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in env_table.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            
            for row_idx, (env_name, env_data) in enumerate(by_env.items(), start=1):
                env_table.rows[row_idx].cells[0].text = env_name
                env_table.rows[row_idx].cells[1].text = str(env_data.get('ok', 0))
                env_table.rows[row_idx].cells[2].text = str(env_data.get('warning', 0))
                env_table.rows[row_idx].cells[3].text = str(env_data.get('critical', 0))
                env_table.rows[row_idx].cells[4].text = str(env_data.get('unknown', 0))
                for cell in env_table.rows[row_idx].cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 3. 카테고리별 점검 결과
        doc.add_heading('3. 카테고리별 점검 결과', level=1)
        
        by_cat = summary.get('by_category', {})
        if by_cat:
            cat_table = doc.add_table(rows=len(by_cat) + 1, cols=5)
            cat_table.style = 'Table Grid'
            
            cat_headers = ['카테고리', '정상', '경고', '위험', '확인불가']
            for i, h in enumerate(cat_headers):
                cat_table.rows[0].cells[i].text = h
                cat_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cat_table.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            
            for row_idx, (cat_name, cat_data) in enumerate(by_cat.items(), start=1):
                cat_table.rows[row_idx].cells[0].text = cat_name
                cat_table.rows[row_idx].cells[1].text = str(cat_data.get('ok', 0))
                cat_table.rows[row_idx].cells[2].text = str(cat_data.get('warning', 0))
                cat_table.rows[row_idx].cells[3].text = str(cat_data.get('critical', 0))
                cat_table.rows[row_idx].cells[4].text = str(cat_data.get('unknown', 0))
                for cell in cat_table.rows[row_idx].cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 4. 상세 점검 결과
        doc.add_heading('4. 상세 점검 결과', level=1)
        
        env_results = {}
        for r in results:
            env = r.get('환경', 'Unknown')
            if env not in env_results:
                env_results[env] = []
            env_results[env].append(r)
        
        section_num = 1
        for env_name, env_data in env_results.items():
            doc.add_heading(f'4.{section_num} {env_name} 환경', level=2)
            
            cat_results = {}
            for r in env_data:
                cat = r.get('카테고리', 'Unknown')
                if cat not in cat_results:
                    cat_results[cat] = []
                cat_results[cat].append(r)
            
            for cat_name, cat_data in cat_results.items():
                doc.add_heading(f'{cat_name}', level=3)
                
                table = doc.add_table(rows=len(cat_data) + 1, cols=5)
                table.style = 'Table Grid'
                
                headers = ['점검ID', '점검항목', '상태', '측정값', '결과메시지']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in table.rows[0].cells[i].paragraphs[0].runs:
                        run.bold = True
                
                for row_idx, r in enumerate(cat_data, start=1):
                    table.rows[row_idx].cells[0].text = r.get('점검ID', '')
                    table.rows[row_idx].cells[1].text = r.get('점검항목', '')
                    
                    status = r.get('상태', '')
                    table.rows[row_idx].cells[2].text = status
                    status_cell = table.rows[row_idx].cells[2]
                    for run in status_cell.paragraphs[0].runs:
                        if status == '정상':
                            run.font.color.rgb = RGBColor(0, 128, 0)
                        elif status == '경고':
                            run.font.color.rgb = RGBColor(255, 165, 0)
                        elif status == '위험':
                            run.font.color.rgb = RGBColor(255, 0, 0)
                        run.bold = True
                    
                    value = r.get('측정값', '')
                    if len(value) > 30:
                        value = value[:30] + "..."
                    table.rows[row_idx].cells[3].text = value
                    table.rows[row_idx].cells[4].text = r.get('결과메시지', '')
                
                doc.add_paragraph()
            
            section_num += 1
        
        # 5. 조치 필요 항목
        issues = [r for r in results if r.get('상태') in ['경고', '위험']]
        if issues:
            doc.add_heading('5. 조치 필요 항목', level=1)
            
            for issue in issues:
                status = issue.get('상태', '')
                icon = "[경고]" if status == '경고' else "[위험]"
                
                para = doc.add_paragraph()
                run = para.add_run(f"{icon} [{issue.get('점검ID')}] {issue.get('점검항목')}")
                run.bold = True
                
                para.add_run(f"\n   - 환경: {issue.get('환경', '')}")
                para.add_run(f"\n   - 대상: {issue.get('점검대상', '')}")
                para.add_run(f"\n   - 상태: {status}")
                para.add_run(f"\n   - 측정값: {issue.get('측정값', '')}")
                para.add_run(f"\n   - 메시지: {issue.get('결과메시지', '')}")
                para.add_run(f"\n   - 중요도: {issue.get('중요도', '')}")
        
        # 6. 서명란
        doc.add_paragraph()
        doc.add_paragraph()
        
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run("점검자: ________________")
        sign_para.add_run("\n\n")
        sign_para.add_run("검토자: ________________")
        sign_para.add_run("\n\n")
        sign_para.add_run(f"점검일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        
        doc.save(filepath)
        return filepath


def generate_reports(results: List[Dict], summary: Dict, config: ReportConfig = None) -> Dict[str, str]:
    """CSV 및 DOCX 보고서 생성"""
    generator = CMPReportGenerator(config)
    generated = {}
    
    csv_path = generator.generate_csv(results, summary)
    generated['csv'] = csv_path
    
    if DOCX_AVAILABLE:
        docx_path = generator.generate_docx(results, summary)
        generated['docx'] = docx_path
    
    return generated


if __name__ == "__main__":
    test_results = [
        {'점검ID': 'OS-001', '점검항목': '디스크 사용량', '카테고리': 'OS', '환경': 'DEV',
         '점검대상': 'master-01', '설명': '테스트', '상태': '정상', '측정값': '45%', 
         '임계치': '80%', '결과메시지': '정상 범위', '중요도': 'high', 
         '점검시간': datetime.now().isoformat()},
    ]
    test_summary = {
        'total': 1, 'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0,
        'by_environment': {'DEV': {'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0}},
        'by_category': {'OS': {'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0}}
    }
    
    config = ReportConfig(company_name="CMP 인프라", team_name="플랫폼팀")
    paths = generate_reports(test_results, test_summary, config)
    print(f"생성된 보고서: {paths}")
